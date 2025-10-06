import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import io

# Not needed anymore

# Настройка страницы и стилей
st.set_page_config(
    page_title='AutoCall — Аналитика',
    layout='wide',
    initial_sidebar_state='expanded'
)

# Кастомные стили CSS
st.markdown("""
<style>
    /* Основной шрифт */
    body {
        font-family: 'Segoe UI', Arial, sans-serif;
        color: #2c3e50;
    }
    
    /* Заголовки */
    h1 {
        color: #2c3e50;
        font-weight: 600;
        font-size: 2.5rem;
        padding-bottom: 1rem;
        border-bottom: 2px solid #3498db;
        margin-bottom: 2rem;
    }
    
    h2 {
        color: #34495e;
        font-weight: 500;
        font-size: 1.8rem;
        margin-top: 2rem;
        margin-bottom: 1rem;
    }

    /* Таблицы */
    .dataframe {
        font-family: 'Segoe UI', Arial, sans-serif !important;
        font-size: 14px !important;
    }
    
    /* Метрики */
    div[data-testid="stMetricValue"] {
        font-size: 2rem !important;
        color: #2980b9 !important;
    }

    /* Подписи к метрикам */
    div[data-testid="stMetricLabel"] {
        font-size: 1rem !important;
        color: #7f8c8d !important;
    }

    /* Отступы для визуального разделения */
    .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
    }

    /* Инфо-блоки */
    div.stAlert {
        font-family: 'Segoe UI', Arial, sans-serif !important;
        border-radius: 8px !important;
        padding: 1rem !important;
    }

    /* Кнопки */
    .stButton > button {
        font-family: 'Segoe UI', Arial, sans-serif !important;
        font-size: 1rem !important;
        padding: 0.5rem 1rem !important;
        border-radius: 8px !important;
    }

    /* Селекторы */
    .stSelectbox > div > div {
        font-family: 'Segoe UI', Arial, sans-serif !important;
    }

    /* Тени для карточек */
    div[data-testid="stMetricValue"] {
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        border-radius: 8px;
        padding: 1rem;
        background: white;
    }

    /* Анимации при наведении */
    .stButton > button:hover {
        transform: translateY(-1px);
        transition: all 0.2s ease;
    }
</style>
""", unsafe_allow_html=True)

# Заголовок с центрированием
st.markdown("<h1 style='text-align: center;'>AutoCall — Аналитика по обзвону пациентов</h1>", unsafe_allow_html=True)

# Описание приложения
st.markdown("""
<div style='text-align: center; padding: 1rem; margin-bottom: 2rem; color: #7f8c8d;'>
    Интерактивный анализ данных обзвона пациентов с визуализацией KPI и CSI по отделениям
</div>
""", unsafe_allow_html=True)


uploaded = st.file_uploader('Загрузите Excel/CSV', type=['xlsx', 'xls', 'csv'], accept_multiple_files=True)
if uploaded:
    def count_answers(row):
        vals = row[question_cols]
        return sum(str(v).strip().isdigit() and 1 <= int(str(v).strip()) <= 10 for v in vals)

    def calculate_csi(row):
        vals = row[question_cols]
        nums = [int(str(v).strip()) for v in vals if str(v).strip().isdigit() and 1 <= int(str(v).strip()) <= 10]
        return sum(nums) / len(nums) if nums else None

    def process_single_file(file):
        """Process a single uploaded file and return processed df."""
        if hasattr(file, 'name') and file.name.lower().endswith('.csv'):
            temp_df = pd.read_csv(file, header=None)
        else:
            temp_df = pd.read_excel(file, header=None)
        temp_df.columns = temp_df.iloc[3]
        temp_df = temp_df.iloc[4:].reset_index(drop=True)
        temp_df = temp_df.dropna(how='all').reset_index(drop=True)
        temp_df.columns = [str(c).strip() for c in temp_df.columns]
        if len(temp_df) > 0 and str(temp_df.iloc[-1,0]).strip().lower().startswith('всего'):
            temp_df = temp_df.iloc[:-1].reset_index(drop=True)

        # Local processing
        question_cols_temp = temp_df.columns[5:13]
        dept_col_temp = temp_df.columns[2]
        temp_df['_answers'] = temp_df.apply(lambda row: sum(str(v).strip().isdigit() and 1 <= int(str(v).strip()) <= 10 for v in row[question_cols_temp]), axis=1)
        temp_df['_csi'] = temp_df.apply(lambda row: sum([int(str(v).strip()) for v in row[question_cols_temp] if str(v).strip().isdigit() and 1 <= int(str(v).strip()) <= 10]) / len([int(str(v).strip()) for v in row[question_cols_temp] if str(v).strip().isdigit() and 1 <= int(str(v).strip()) <= 10]) if [int(str(v).strip()) for v in row[question_cols_temp] if str(v).strip().isdigit() and 1 <= int(str(v).strip()) <= 10] else None, axis=1)
        temp_df['month'] = file.name
        return temp_df, question_cols_temp, dept_col_temp

    def compute_dept_stats(temp_df, question_cols_temp, dept_col_temp):
        """Compute department stats for a df."""
        num_col = temp_df.columns[0]
        dept_stats = temp_df.groupby(dept_col_temp).agg(
            звонков=(num_col, 'count'),
            средний_CSI=('_csi', 'mean'),
            ответили_все=(num_col, lambda x: (temp_df.loc[x.index, '_answers'] == len(question_cols_temp)).sum()),
            ответили_хотябы=(num_col, lambda x: (temp_df.loc[x.index, '_answers'] > 0).sum()),
            количество_CSI=(num_col, lambda x: temp_df.loc[x.index, '_csi'].notna().sum()),
        )
        dept_stats['%_ответили_все'] = dept_stats['ответили_все'] / dept_stats['звонков'] * 100
        dept_stats['%_ответили_хотябы'] = dept_stats['ответили_хотябы'] / dept_stats['звонков'] * 100
        dept_stats['средний_CSI'] = dept_stats['средний_CSI'].round(1)
        return dept_stats

    try:
        df_list = []
        for file in uploaded:
            temp_df, question_cols_temp, dept_col_temp = process_single_file(file)
            df_list.append(temp_df)
        df = pd.concat(df_list, ignore_index=True)
    except Exception as e:
        st.error('Ошибка чтения файла: ' + str(e))
    else:
        st.write('Размер:', df.shape)
        # Скрывающаяся секция с исходными данными
        with st.expander("📋 Исходные данные", expanded=False):
            st.dataframe(df.head(50), width='stretch')

        question_cols = df.columns[5:13]
        dept_col = df.columns[2]
        total_calls = len(df)

        def count_answers(row):
            vals = row[question_cols]
            return sum(str(v).strip().isdigit() and 1 <= int(str(v).strip()) <= 10 for v in vals)

        def calculate_csi(row):
            vals = row[question_cols]
            nums = [int(str(v).strip()) for v in vals if str(v).strip().isdigit() and 1 <= int(str(v).strip()) <= 10]
            return sum(nums) / len(nums) if nums else None

        # The df already has '_answers' and '_csi' from the process_single_file

        all_answered = (df['_answers'] == len(question_cols)).sum()
        any_answered = (df['_answers'] > 0).sum()
        percent_all = all_answered / total_calls * 100 if total_calls else 0
        percent_any = any_answered / total_calls * 100 if total_calls else 0

        # Среднее количество ответов по пациентам с хотя бы одним ответом
        avg_answers_with_some = np.round(np.mean(df['_answers'][df['_answers'] > 0]), 1) if any_answered else None

        # Средний CSI по пациентам с ответами
        valid_csi = df['_csi'].dropna()
        avg_csi = np.round(np.mean(valid_csi), 1) if len(valid_csi) else None

        st.markdown("## 📊 Ключевые показатели")

        # KPI метрики в пять колонок
        col1, col2, col3, col4, col5 = st.columns(5)
        with col1:
            st.metric(
                'Всего обзвоненных пациентов',
                f"{total_calls:,}",
                help="Общее количество пациентов в выгрузке"
            )
        with col2:
            st.metric(
                'Ответили на все вопросы',
                f"{all_answered:,}",
                delta=f"{percent_all:.1f}%",
                help="Количество и процент пациентов, ответивших на все вопросы"
            )
        with col3:
            st.metric(
                'Ответили хотя бы на один вопрос',
                f"{any_answered:,}",
                delta=f"{percent_any:.1f}%",
                help="Количество и процент пациентов, ответивших хотя бы на один вопрос"
            )
        with col4:
            st.metric(
                'Среднее кол-во ответов',
                f"{avg_answers_with_some:.1f}" if avg_answers_with_some else "Нет данных",
                help="Среднее количество ответов среди пациентов, ответивших хотя бы на один вопрос"
            )
        with col5:
            st.metric(
                'Средний CSI',
                f"{avg_csi:.1f}" if avg_csi else "Нет данных",
                help="Среднее значение CSI среди пациентов, ответивших хотя бы на один вопрос"
            )

        st.markdown("## 📈 CSI по отделениям")

        dept_stats = compute_dept_stats(df, question_cols, dept_col)

        # Форматирование таблицы с градиентной подсветкой
        styled_stats = dept_stats.reset_index().style\
            .background_gradient(subset=['средний_CSI'], cmap='RdYlGn')\
            .format({
                'средний_CSI': '{:.1f}',
                '%_ответили_все': '{:.1f}%',
                '%_ответили_хотябы': '{:.1f}%'
            })\
            .set_properties(**{
                'font-size': '14px',
                'font-family': 'Segoe UI, Arial, sans-serif',
                'text-align': 'center'
            })
        
        st.dataframe(styled_stats, width='stretch')

        st.markdown("## 📊 Сравнение отделений по вопросам")

        # Выбор вопросов для отображения
        selected_question = st.selectbox(
            'Выберите вопрос для анализа:',
            options=list(enumerate(question_cols, 1)),
            index=0,
            format_func=lambda x: f'Вопрос {x[0]}: {x[1]}'
        )

        # График для выбранного вопроса
        i, qcol = selected_question
        q_stats = df.groupby(dept_col)[qcol].apply(
            lambda vals: np.mean(lst) if (lst := [int(str(v).strip()) for v in vals if str(v).strip().isdigit() and 1 <= int(str(v).strip()) <= 10]) else np.nan
        )
        q_counts = df.dropna(subset=[qcol]).groupby(dept_col)[qcol].count()

        fig = go.Figure()
        fig.add_bar(
            x=q_stats.index,
            y=q_stats.values,
            marker_color='skyblue',
            text=q_stats.values.round(1),
            textposition='auto',
            textfont=dict(size=18),
            customdata=q_counts.reindex(q_stats.index),
            hovertemplate="Средний балл: %{y:.1f}<br>Количество ответов: %{customdata}",
        )

        # Настройка графика
        fig.update_layout(
            height=450,
            margin=dict(t=100, b=50, l=50, r=50),
            yaxis_range=[0, 10],
            title=f'<b>Средний балл по вопросу {i}</b><br>{qcol}',
            xaxis_title='Отделение',
            yaxis_title='Средний балл',
            font=dict(size=14),
            xaxis=dict(tickfont=dict(size=14)),
            yaxis=dict(tickfont=dict(size=14)),
            title_font=dict(size=18)
        )

        # Вывод с минимальной конфигурацией
        st.plotly_chart(fig, use_container_width=True)

        # Сравнение периодов
        if len(uploaded) >= 2:
            st.markdown("## 📊 Сравнение периодов")
            files = [f.name for f in uploaded]
            col_comp1, col_comp2 = st.columns(2)
            with col_comp1:
                file1 = st.selectbox("Выберите первый файл", files, index=0, key='file1')
            with col_comp2:
                file2 = st.selectbox("Выберите второй файл", files, index=min(1, len(files)-1), key='file2')

            if file1 != file2:
                df1, question_cols1, dept_col1 = process_single_file(uploaded[files.index(file1)])
                df2, question_cols2, dept_col2 = process_single_file(uploaded[files.index(file2)])
                dept_stats1 = compute_dept_stats(df1, question_cols1, dept_col1)
                dept_stats2 = compute_dept_stats(df2, question_cols2, dept_col2)

                st.markdown(f"### Сравнение {file1} и {file2}")
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown(f"**{file1}**")
                    st.dataframe(dept_stats1.reset_index(), width='stretch')
                with col2:
                    st.markdown(f"**{file2}**")
                    st.dataframe(dept_stats2.reset_index(), width='stretch')

                # Comparison chart
                common_depts = set(dept_stats1.index) & set(dept_stats2.index)
                if common_depts:
                    comp_df = pd.DataFrame(index=list(common_depts))
                    comp_df[file1] = dept_stats1.loc[list(common_depts), 'средний_CSI']
                    comp_df[file2] = dept_stats2.loc[list(common_depts), 'средний_CSI']

                    fig_comp = go.Figure()
                    for col in comp_df.columns:
                        counts = dept_stats1['количество_CSI'][list(common_depts)] if col == file1 else dept_stats2['количество_CSI'][list(common_depts)]
                        fig_comp.add_bar(
                            name=col,
                            x=comp_df.index,
                            y=comp_df[col],
                            text=comp_df[col].round(1),
                            textposition='auto',
                            textfont=dict(size=18),
                            customdata=counts,
                            hovertemplate="Средний CSI: %{y:.1f}<br>Количество CSI: %{customdata}",
                        )
                    fig_comp.update_layout(
                        barmode='group',
                        title="Сравнение среднего CSI по отделениям",
                        xaxis_title='Отделение',
                        yaxis_title='Средний CSI',
                        font=dict(size=14),
                        margin=dict(t=100, b=50, l=50, r=50),
                        title_font=dict(size=16)
                    )
                    st.plotly_chart(fig_comp, use_container_width=True)
                else:
                    st.write("Нет общих отделений для сравнения.")

        # Экспорт данных
        st.markdown("## 💾 Экспорт данных")

        # Generate Word document report
        try:
            from docx import Document
            from docx.shared import Pt

            doc = Document()

            # Set default font for the document
            from docx.oxml.shared import OxmlElement, nsid
            font_style = doc.styles['Normal'].font
            font_style.name = 'Calibri'

            doc.add_heading('AutoCall Аналитика - Отчет', 0)
            doc.add_paragraph('.')
            para1 = doc.add_paragraph(f'Всего обзвоненных пациентов: {total_calls}')
            for run in para1.runs:
                run.font.name = 'Calibri'
            para2 = doc.add_paragraph(f'Ответили на все вопросы: {percent_all:.1f}%')
            for run in para2.runs:
                run.font.name = 'Calibri'
            para3 = doc.add_paragraph(f'Ответили хотя бы на один вопрос: {percent_any:.1f}%')
            for run in para3.runs:
                run.font.name = 'Calibri'
            para4 = doc.add_paragraph(f'Среднее кол-во ответов: {avg_answers_with_some:.1f}')
            for run in para4.runs:
                run.font.name = 'Calibri'
            para5 = doc.add_paragraph(f'Средний CSI: {avg_csi:.1f}')
            for run in para5.runs:
                run.font.name = 'Calibri'
            doc.add_paragraph('.')
            doc.add_heading('Статистика по отделениям', level=2)
            table = doc.add_table(rows=1, cols=dept_stats.shape[1] + 1)
            table.style = 'Table Grid'

            # Set font for all table cells to Calibri
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Calibri'
                            run.font.size = Pt(11)

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Отделение'
            for i, col in enumerate(dept_stats.columns):
                hdr_cells[i+1].text = col
            for idx, (dept, row) in enumerate(dept_stats.iterrows()):
                row_cells = table.add_row().cells
                row_cells[0].text = dept
                for i, val in enumerate(row):
                    if isinstance(val, float):
                        row_cells[i+1].text = f"{val:.1f}"
                    else:
                        row_cells[i+1].text = str(val)
            buffer = io.BytesIO()
            doc.save(buffer)
            pdf_file_data = buffer.getvalue()
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            fname = "report.docx"
        except ImportError:
            pdf_file_data = "Установите python-docx для генерации документа Word: pip install python-docx\n\n"
            pdf_file_data += f"Всего обзвоненных пациентов: {total_calls}\n"
            pdf_file_data += f"Ответили на все вопросы: {percent_all:.1f}%\n"
            pdf_file_data += f"Ответили хотя бы на один вопрос: {percent_any:.1f}%\n"
            pdf_file_data += f"Среднее кол-во ответов: {avg_answers_with_some:.1f}\n"
            pdf_file_data += f"Средний CSI: {avg_csi:.1f}\n\n"
            pdf_file_data += "Статистика по отделением:\n"
            pdf_file_data += dept_stats.reset_index().to_csv(index=False, encoding='utf-8-sig')
            pdf_file_data = pdf_file_data.encode('utf-8-sig')
            mime = "text/plain"
            fname = "report.txt"

        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                "⬇️ Скачать очищенные данные (CSV)",
                df.to_csv(index=False, encoding='utf-8-sig'),
                "cleaned_data.csv",
                "text/csv",
                key='download-csv'
            )
        with col2:
            st.download_button(
                "⬇️ Скачать статистику по отделениям (CSV)",
                dept_stats.reset_index().to_csv(index=False, encoding='utf-8-sig'),
                "department_stats.csv",
                "text/csv",
                key='download-stats'
            )
